# Document Converter Module
import io
import subprocess
import os
from pypdf import PdfReader
import docx
from pdf2image import convert_from_bytes
import zipfile
from pathlib import Path
from loguru import logger

# Ensure logs directory exists and use it for logging (using pathlib)
BASE_DIR = Path(__file__).resolve().parent.parent
LOG_DIR = BASE_DIR / "logs"
LOG_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE = LOG_DIR / "docconv.log"

# Add a handler to the existing logger for this module only
logger.add(LOG_FILE, rotation="10 MB", retention="1 day", filter=lambda record: record["extra"].get("module") == "documentconvert")

# Create a module-specific logger instead of using the global one
doc_logger = logger.bind(module="documentconvert")

SUPPORT_PDF_IMAGE = ['jpeg', 'png', 'tiff', 'webp', 'bmp']

def libreoffice_convert(input_bytes, input_ext, output_ext):
    import tempfile

    # Only allow safe extensions
    allowed_exts = {"pdf", "docx", "txt", "odt", "rtf"}
    if input_ext.lower() not in allowed_exts or output_ext.lower() not in allowed_exts | {"pdf"}:
        raise ValueError("Invalid file extension for conversion.")

    # Write input to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{input_ext}") as tmp_in:
        tmp_in.write(input_bytes)
        tmp_in.flush()
        input_path = Path(tmp_in.name).resolve()
    # Output will be in the same dir
    out_dir = input_path.parent
    try:
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", output_ext,
            "--outdir", str(out_dir),
            str(input_path)
        ]
        # Security: All command arguments are validated
        # - output_ext and input_ext are whitelisted.
        # - out_dir and input_path are from secure temp files.
        # - No user-controlled input is passed unsanitized.
        # - shell=False is enforced.
        subprocess.run(
            cmd,
            check=True,
            shell=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )
        output_path = input_path.with_suffix(f".{output_ext}").resolve()
        # Validate output path is in the expected directory
        if not str(output_path).startswith(str(out_dir)):
            raise RuntimeError("Unexpected output path from libreoffice.")
        with open(output_path, "rb") as f:
            result_bytes = f.read()
        os.remove(output_path)
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice conversion failed: {e.stderr.decode()}")
        raise
    except Exception as e:
        print(f"Error during LibreOffice conversion: {e}")
        raise
    finally:
        os.remove(str(input_path))
    return result_bytes

def pdf_to_image_bytes(pdf_bytes, image_format, settings=None):
    """
    Convert PDF bytes to image bytes in the specified format.
    Returns a list of (filename, image bytes) tuples (one per page).
    settings: dict, may include 'dpi' (int), 'quality' (int, for JPEG/WEBP), etc.
        - 'dpi': int (default 200) - DPI for image rendering.
        - 'quality': int (for JPEG/WEBP only) - quality (1-100).
    """
    dpi = settings.get('dpi', 200) if settings else 200
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    img_bytes_list = []
    for idx, img in enumerate(images):
        img_byte_arr = io.BytesIO()
        save_kwargs = {}
        if image_format.lower() in ('jpeg', 'webp') and settings and 'quality' in settings:
            save_kwargs['quality'] = settings['quality']
        img.save(img_byte_arr, format=image_format.upper(), **save_kwargs)
        filename = f"page_{idx+1}.{image_format.lower()}"
        img_bytes_list.append((filename, img_byte_arr.getvalue()))
    return img_bytes_list

def _pdf_to_txt(input_stream):
    pdf = PdfReader(input_stream)
    text = ""
    for page in pdf.pages:
        text += page.extract_text()
    return text.encode('utf-8')

def _txt_to_pdf(input_stream):
    doc = docx.Document()
    text = input_stream.read().decode('utf-8')
    for line in text.split('\n'):
        doc.add_paragraph(line)
    tmp_bytes = io.BytesIO()
    doc.save(tmp_bytes)
    pdf_bytes = libreoffice_convert(tmp_bytes.getvalue(), "docx", "pdf")
    return pdf_bytes

def _docx_to_txt(input_stream):
    doc = docx.Document(input_stream)
    text = ""
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text.encode('utf-8')

def _txt_to_docx(input_stream):
    doc = docx.Document()
    text = input_stream.read().decode('utf-8')
    for line in text.split('\n'):
        doc.add_paragraph(line)
    tmp_bytes = io.BytesIO()
    doc.save(tmp_bytes)
    return tmp_bytes.getvalue()

def _pdf_to_docx(contents):
    return libreoffice_convert(contents, "pdf", "docx")

def _pdf_to_image_zip(contents, to_format, settings):
    image_bytes_list = pdf_to_image_bytes(contents, to_format, settings)
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for filename, img_bytes in image_bytes_list:
            zipf.writestr(filename, img_bytes)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def _docx_to_pdf(contents):
    return libreoffice_convert(contents, "docx", "pdf")

def convert_document(contents, from_format, to_format, settings=None):
    """
    Convert document from one format to another.
    Currently supports limited conversions between PDF, DOCX, and TXT.
    Only PDF to image supports settings.
    """
    from_format = from_format.lower()
    to_format = to_format.lower()
    input_stream = io.BytesIO(contents)
    output_stream = io.BytesIO()

    print(f"Converting from {from_format} to {to_format}")

    if from_format == 'pdf' and to_format == 'txt':
        output_stream.write(_pdf_to_txt(input_stream))
    elif from_format == 'txt' and to_format == 'pdf':
        output_stream.write(_txt_to_pdf(input_stream))
    elif from_format == 'docx' and to_format == 'txt':
        output_stream.write(_docx_to_txt(input_stream))
    elif from_format == 'txt' and to_format == 'docx':
        output_stream.write(_txt_to_docx(input_stream))
    elif from_format == 'pdf' and to_format == 'docx':
        output_stream.write(_pdf_to_docx(contents))
    elif from_format == 'pdf' and to_format in SUPPORT_PDF_IMAGE:
        return _pdf_to_image_zip(contents, to_format, settings)
    elif from_format == 'docx' and to_format == 'pdf':
        output_stream.write(_docx_to_pdf(contents))
    else:
        raise ValueError(f"Conversion from {from_format} to {to_format} not supported yet.")

    return output_stream.getvalue()

